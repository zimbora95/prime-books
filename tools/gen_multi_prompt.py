#!/usr/bin/env python
"""Generate PROMPT PDF TO BOOK MULTIPLE SUBJECTS - ALL.docx with all 27 prompts (Y1-Y3 x 9 subjects).
Each subject prompt includes the QR code instruction at the end of the 'Remake the ENTIRE book' paragraph.
"""
import os
from docx import Document
from docx.shared import Pt

BASE = r"H:\Shared drives\Prime Books"
LOCAL_BASE = r"C:\Users\alexa\Documents\The Prime Books"

# folder name per year per subject
FOLDER_MAP = {
    "Year 01": {
        "Art & Design": "Art & Design",
        "Computing & Robotics": "Computing & Robotics",
        "English": "English",
        "Global Perspectives": "Global Perspectives",
        "Mathematics": "Mathematics",
        "Music & Drama": "Music & Drama",
        "Physical Education": "Physical Education",
        "Portuguese": "Portuguese 1st",
        "Science": "Science",
    },
    "Year 02": {
        "Art & Design": "Art & Design",
        "Computing & Robotics": "Computing & Robotics",
        "English": "English",
        "Global Perspectives": "Global Perspectives",
        "Mathematics": "Mathematics",
        "Music & Drama": "Music & Acting",
        "Physical Education": "Physical Education - Gym",
        "Portuguese": "Portuguese 1st",
        "Science": "Science & Lab",
    },
    "Year 03": {
        "Art & Design": "Art & Design",
        "Computing & Robotics": "Computing & Robotics",
        "English": "English",
        "Global Perspectives": "Global Perspectives",
        "Mathematics": "Mathematics",
        "Music & Drama": "Music & Acting",
        "Physical Education": "Physical Education - Gym",
        "Portuguese": "Portuguese 1st",
        "Science": "Science & Lab",
    },
}

# primary input PDF per subject per year (filename only, inside PDF/Input/)
INPUT_PDFS = {
    ("Year 01", "Art & Design"): "Y1 Art & Design.xlsx",
    ("Year 01", "Computing & Robotics"): "Computing Y01 - Scheme of Work.xlsx",
    ("Year 01", "English"): "English Y01 - Student Book.pdf",
    ("Year 01", "Global Perspectives"): "Y1 Global Perspectives.xlsx",
    ("Year 01", "Mathematics"): "Mathematics Year 01 - Student Book.pdf",
    ("Year 01", "Music & Drama"): "Music and Drama Y01 - Scheme of Work.xlsx",
    ("Year 01", "Physical Education"): "Year 1 Gym Prime Books.pdf",
    ("Year 01", "Portuguese"): "Portugues-Y1-Com-Bonecos (1 Vf).pdf",
    ("Year 01", "Science"): "cambridge-primary-science-student-book.pdf",
    ("Year 02", "Art & Design"): "Year 2 Art & Design.xlsx",
    ("Year 02", "Computing & Robotics"): "PrimeBooks-Computing-Year-2-Student.pdf",
    ("Year 02", "English"): "Cópia de English Y02 - Student Book.pdf",
    ("Year 02", "Global Perspectives"): "PrimeBooks-Global-Year-2-Student.pdf",
    ("Year 02", "Mathematics"): "Collins Mathematics Y02 - Student Book.pdf",
    ("Year 02", "Music & Drama"): "Edu360 Topic (edu360.topic) Music Y2.xlsx",
    ("Year 02", "Physical Education"): "Year 2 Gym Prime Books.pdf",
    ("Year 02", "Portuguese"): "Portugues-Y2-Com-Bonecos (1)vf.pdf",
    ("Year 02", "Science"): None,  # no input PDF found
    ("Year 03", "Art & Design"): "Year 3_Art & Design_2025-2026_Scheme of Work.xlsx",
    ("Year 03", "Computing & Robotics"): "Year-3-Computing-Prime-School.pdf",
    ("Year 03", "English"): None,  # no input found
    ("Year 03", "Global Perspectives"): "PrimeBooks-Global-Perspectives-Year-3-PRINT.pdf",
    ("Year 03", "Mathematics"): "Mathematics Y03 - Student Book.pdf",
    ("Year 03", "Music & Drama"): "Edu360 Topic (edu360.topic) Music Y3.xlsx",
    ("Year 03", "Physical Education"): "Year 3 Gym Prime Books.pdf",
    ("Year 03", "Portuguese"): "Portugues-Y3-Com-Bonecos.pdf",
    ("Year 03", "Science"): "Collins Science Student Book 3.pdf",
}

# cover template / output PDF per subject per year
OUTPUT_PDFS = {
    ("Year 01", "Art & Design"): "Prime Book - Art & Design - Year 1 - Student Book.pdf",
    ("Year 01", "Computing & Robotics"): "Prime Book - Computing & Robotics - Year 1 - Student Book.pdf",
    ("Year 01", "English"): "Prime Book - English - Year 1 - Student Book.pdf",
    ("Year 01", "Global Perspectives"): "Prime Book - Global Perspectives - Year 1 - Student Book.pdf",
    ("Year 01", "Mathematics"): "Prime Books Mathematics Year 01 - Student Book.pdf",
    ("Year 01", "Music & Drama"): "Prime Books Music and Acting Year 01 - Student Manual.pdf",
    ("Year 01", "Physical Education"): "COVER TEMPLATE - Physical Education & Gym - Year 1.pdf",
    ("Year 01", "Portuguese"): "Prime Book - Portuguese 1st - Year 1 - Student Book.pdf",
    ("Year 01", "Science"): "Prime Books Science and Lab Year 01 - Student Book.pdf",
    ("Year 02", "Art & Design"): "COVER TEMPLATE - Art & Design - Year 2.pdf",
    ("Year 02", "Computing & Robotics"): "COVER TEMPLATE - Computing & Robotics - Year 2.pdf",
    ("Year 02", "English"): "Prime Book - English - Year 2 - Student Book.pdf",
    ("Year 02", "Global Perspectives"): "COVER TEMPLATE - Global Perspectives - Year 2.pdf",
    ("Year 02", "Mathematics"): "Prime Book - Mathematics - Year 2 - Student Book.pdf",
    ("Year 02", "Music & Drama"): "COVER TEMPLATE - Music & Acting - Year 2.pdf",
    ("Year 02", "Physical Education"): "COVER TEMPLATE - Physical Education & Gym - Year 2.pdf",
    ("Year 02", "Portuguese"): "Prime Book - Portuguese 1st - Year 2 - Student Book.pdf",
    ("Year 02", "Science"): "COVER TEMPLATE - Science & Lab - Year 2.pdf",
    ("Year 03", "Art & Design"): "COVER TEMPLATE - Art & Design - Year 3.pdf",
    ("Year 03", "Computing & Robotics"): "COVER TEMPLATE - Computing & Robotics - Year 3.pdf",
    ("Year 03", "English"): "COVER TEMPLATE - English - Year 3.pdf",
    ("Year 03", "Global Perspectives"): "COVER TEMPLATE - Global Perspectives - Year 3.pdf",
    ("Year 03", "Mathematics"): "Prime Book - Mathematics - Year 3 - Student Book.pdf",
    ("Year 03", "Music & Drama"): "COVER TEMPLATE - Music & Acting - Year 3.pdf",
    ("Year 03", "Physical Education"): "COVER TEMPLATE - Physical Education & Gym - Year 3.pdf",
    ("Year 03", "Portuguese"): "COVER TEMPLATE - Portuguese 1st - Year 3.pdf",
    ("Year 03", "Science"): "COVER TEMPLATE - Science & Lab - Year 3.pdf",
}

YEAR_LABEL = {"Year 01": "Year 1", "Year 02": "Year 2", "Year 03": "Year 3"}
YEAR_ZERO = {"Year 01": "01", "Year 02": "02", "Year 03": "03"}
YEAR_LOCAL = {"Year 01": "Year01", "Year 02": "Year02", "Year 03": "Year03"}

SUBJECTS = [
    "Art & Design",
    "Computing & Robotics",
    "English",
    "Global Perspectives",
    "Mathematics",
    "Music & Drama",
    "Physical Education",
    "Portuguese",
    "Science",
]

QR_INSTRUCTION = (
    "Every unit should intelligently include QR codes that extend learning beyond the printed page. "
    "These are not decorative elements. Every QR code must have a clear educational purpose and "
    "provide genuine value."
)

ROLE_TEXT = (
    'You are a world-class book designer + senior front-end engineer producing COMPLETE, print-ready\n'
    'educational books for Prime School (www.primeschool.pt), an international Cambridge school. '
    'Awwwards-tier quality is the bar. No lorem ipsum, no placeholders, no \u201cai slop\u201d, no "TODO", '
    'no "sample pages". Fully finished means FULLY FINISHED. every sentence rewritten, every image '
    'generated, every page fitting its frame, every claim verified\n'
    'with real tool output. You have carte blanche on creative decisions.'
)

REMAKE_TEXT = (
    'Remake the ENTIRE book based on this "{input_path}" as an original Prime School publication: '
    'professional cover, imprint/copyright page, contents, introduction, how-to-use, getting-set-up, '
    'etc all units complete with every feature (For example: Get started, Scenario, Learning outcomes, '
    'Warm up, Do you remember, Learn, Practise, Did you\n'
    'know, Go further, Challenge yourself, Final project, Evaluation, What can you do, Keywords), '
    'a full glossary, etc and a back cover. This is a real student textbook, not a sample. '
    + QR_INSTRUCTION
)


def build_prompt_subject(doc, year_folder, subject):
    ynum = YEAR_LABEL[year_folder]
    yzero = YEAR_ZERO[year_folder]
    ylocal = YEAR_LOCAL[year_folder]
    actual_folder = FOLDER_MAP[year_folder][subject]
    subj_path = os.path.join(BASE, year_folder, actual_folder)
    input_dir = os.path.join(subj_path, "PDF", "Input")
    output_dir = os.path.join(subj_path, "PDF", "Output")
    images_dir = os.path.join(subj_path, "IMAGES")
    local_dir = os.path.join(LOCAL_BASE, f"{ylocal}-{actual_folder}")

    # input PDF
    input_file = INPUT_PDFS.get((year_folder, subject))
    if input_file:
        input_path = os.path.join(input_dir, input_file)
    else:
        input_path = os.path.join(input_dir, "(provide input PDF here)")

    # output PDF / cover
    output_file = OUTPUT_PDFS.get((year_folder, subject), "(provide output PDF here)")
    output_path = os.path.join(output_dir, output_file)

    # subject header
    doc.add_paragraph(f"YEAR {ynum.split()[-1]} - {subject.upper()}")
    doc.add_paragraph("")

    # ROLE
    doc.add_paragraph("# ROLE")
    doc.add_paragraph("")
    doc.add_paragraph(ROLE_TEXT)
    doc.add_paragraph("")

    # CONTEXT & ASSETS
    doc.add_paragraph("# CONTEXT & ASSETS")
    doc.add_paragraph("")
    doc.add_paragraph("INPUT .PDF PATH")
    doc.add_paragraph("")
    doc.add_paragraph(f'"{input_path}"')
    doc.add_paragraph("")

    doc.add_paragraph("OUTPUT .PDF PATH ")
    doc.add_paragraph("")
    cover_line = (
        f'"{output_path}"'
    )
    if "COVER TEMPLATE" in output_file:
        cover_line += (
            f' at the moment thats the bookcover (page 1) say substitute that as the book cover '
            f'and put on this folder whole prime book at {output_dir}'
        )
    else:
        cover_line += (
            f' put on this folder whole prime book at {output_dir}'
        )
    doc.add_paragraph(cover_line)
    doc.add_paragraph("")

    doc.add_paragraph("OUTPUT .MD PATH")
    doc.add_paragraph("")
    md_line = (
        f'OUTPUT at {os.path.join(output_dir, "MD")} ALSO A FULL STRUCTURED .MD (MARKDOWN)'
        'TO BE EASIER TO MAKE CHANGES IN THE FUTURE IF NEEDED \n'
        'OR OTHER FILES YOU NEED TO FOR THIS TASK. '
    )
    doc.add_paragraph(md_line)
    doc.add_paragraph("")

    doc.add_paragraph("OUTPUT IMAGES PATH")
    doc.add_paragraph("")
    doc.add_paragraph(images_dir)
    doc.add_paragraph("")

    doc.add_paragraph("Prime Books & Prime School logo ")
    doc.add_paragraph("")
    doc.add_paragraph(os.path.join(BASE, "Logo"))
    doc.add_paragraph("")

    doc.add_paragraph(
        'If you need to generate images use openrouter.ai, model openai/gpt-image-2 and output images '
        'with quality set at medium. To prevent the \u201ccutted\u201d images on \u201cbanner images\u201d, '
        'generate images horizontally to have the entire image generated inside it, etc.'
    )
    doc.add_paragraph("")

    doc.add_paragraph("For Internet search you can use Tavily, key already in Hermes config.")
    doc.add_paragraph("")

    # Remake paragraph with QR instruction
    remake = REMAKE_TEXT.format(input_path=input_path)
    doc.add_paragraph(remake)
    doc.add_paragraph("")

    doc.add_paragraph("")
    doc.add_paragraph(
        '- Write an original imprint page \u00a9 Prime School 2026, with an independent-publication '
        'notice, trademark acknowledgements, etc'
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        '- British English. NO em-dashes anywhere (use commas or colons; en-dashes only for ranges '
        'like Ages 11\u201314).'
    )
    doc.add_paragraph("")

    # DESIGN SYSTEM
    doc.add_paragraph("# DESIGN SYSTEM")
    doc.add_paragraph("")
    doc.add_paragraph("Invent one coherent premium top tier, age appropriate")
    doc.add_paragraph("")
    design_line = (
        f'As a top tier pro you have \u201ccart\u00e9 blanche\u201d to do this task and enhance details '
        f'as you see fit. You can spawn sub-agents if you need, it\u2019s all up to you. '
        f'Create Prime Book {subject} Year {yzero} - Student Book. '
        f'If you want to work locally you can here "{local_dir}" '
        f'then sync with {subj_path}'
    )
    doc.add_paragraph(design_line)
    doc.add_paragraph("")
    doc.add_paragraph("")


def main():
    doc = Document()
    # Title
    doc.add_paragraph("PROMPT PDF TO BOOK MULTIPLE SUBJECTS - ALL YEARS")
    doc.add_paragraph("")

    for year_folder in ["Year 01", "Year 02", "Year 03"]:
        for subject in SUBJECTS:
            build_prompt_subject(doc, year_folder, subject)

    out_path = os.path.join(BASE, "PROMPT PDF TO BOOK MULTIPLE SUBJECTS - ALL.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
